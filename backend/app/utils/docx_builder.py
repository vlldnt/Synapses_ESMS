import base64
import re
from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _parse_markdown(text):
    """Parse markdown text into a list of block dicts."""
    blocks = []
    lines = (text or '').split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith('### '):
            blocks.append({'type': 'h3', 'text': stripped[4:]})
        elif stripped.startswith('## '):
            blocks.append({'type': 'h2', 'text': stripped[3:]})
        elif stripped.startswith('# '):
            blocks.append({'type': 'h1', 'text': stripped[2:]})
        elif re.match(r'^-{3,}$', stripped):
            blocks.append({'type': 'hr'})
        elif stripped.startswith('- ') or stripped.startswith('* '):
            blocks.append({'type': 'bullet', 'text': stripped[2:]})
        elif stripped == '':
            blocks.append({'type': 'empty'})
        else:
            blocks.append({'type': 'paragraph', 'text': stripped})

        i += 1
    return blocks


def _add_inline(paragraph, text):
    """Add text with **bold** support."""
    parts = re.split(r'(\*\*.+?\*\*)', text or '')
    for part in parts:
        m = re.match(r'^\*\*(.+)\*\*$', part)
        if m:
            run = paragraph.add_run(m.group(1))
            run.bold = True
        else:
            paragraph.add_run(part)


def generate_docx_base64(text, child_name=None, educator_name=None, educator_role=None):
    """Convert markdown text to DOCX and return base64 string."""
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    # Header
    if child_name or educator_name:
        header = doc.sections[0].header
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        header_para.text = ''
        parts = []
        if child_name:
            parts.append(f'Référence : {child_name}')
        if educator_name:
            label = educator_name
            if educator_role:
                label += f' - {educator_role}'
            parts.append(f'Professionnel : {label}')
        run = header_para.add_run(' | '.join(parts))
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    BRAND = RGBColor(0x0D, 0x66, 0xD4)

    blocks = _parse_markdown(text)
    for block in blocks:
        t = block['type']

        if t == 'h1':
            p = doc.add_paragraph()
            run = p.add_run(block['text'])
            run.bold = True
            run.font.size = Pt(16)
            run.font.color.rgb = BRAND
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(4)

        elif t == 'h2':
            p = doc.add_paragraph()
            run = p.add_run(block['text'])
            run.bold = True
            run.font.size = Pt(13)
            run.font.color.rgb = BRAND
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(2)

        elif t == 'h3':
            p = doc.add_paragraph()
            run = p.add_run(block['text'])
            run.bold = True
            run.font.size = Pt(11)
            p.paragraph_format.space_before = Pt(6)

        elif t == 'bullet':
            p = doc.add_paragraph(style='List Bullet')
            _add_inline(p, block['text'])
            p.paragraph_format.space_after = Pt(2)

        elif t == 'hr':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)

        elif t == 'empty':
            doc.add_paragraph()

        else:
            p = doc.add_paragraph()
            _add_inline(p, block['text'])
            p.paragraph_format.space_after = Pt(3)

    buf = BytesIO()
    doc.save(buf)
    return base64.b64encode(buf.getvalue()).decode('utf-8')
